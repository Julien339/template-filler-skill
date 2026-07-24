"""
Opens the ORIGINAL .docx (the same file extract_docx.py read) and writes new
text into specific runs identified by ID in changes.json, then saves the
result to <output.docx>. Same safety rationale as apply_pptx.py -- ID-only
addressing, in-place run.text mutation (formatting untouched), never a
global search-and-replace.

changes.json shape: [{"id": "3/0", "new_text": "Stellantis"}, ...]

An ID with no trailing run index (e.g. "3" instead of "3/0") is a
PARAGRAPH-level replacement: the new text goes into the paragraph's first
run and every other run in that paragraph is blanked -- see extract_docx.py's
docstring for why a paragraph can have more runs than are meaningful.

Usage:
    python apply_docx.py <template.docx> <changes.json> <output.docx>
"""
import sys
import json
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def replace_paragraph_text(paragraph, new_text):
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return
    runs[0].text = new_text
    for extra in runs[1:]:
        extra.text = ""


def set_run_text(paragraph, run_idx, new_text):
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
    elif run_idx >= len(runs):
        runs[-1].text = new_text
    else:
        runs[run_idx].text = new_text


def resolve_paragraph(blocks, doc, path_parts):
    """path_parts is the change id split on '/', e.g. ["3"], ["3","0"],
    ["5","r1c2","0"], ["5","r1c2","0","0"], ["h0","0"], ["h0","0","0"]."""
    head = path_parts[0]

    if head[0] in ("h", "f"):
        section_idx = int(head[1:])
        section = doc.sections[section_idx]
        part = section.header if head[0] == "h" else section.footer
        para_idx = int(path_parts[1])
        return part.paragraphs[para_idx], path_parts[2:]

    block = blocks[int(head)]
    rest = path_parts[1:]

    if isinstance(block, Table):
        row_s, col_s = rest[0][1:].split("c")
        cell = block.cell(int(row_s), int(col_s))
        para_idx = int(rest[1])
        return cell.paragraphs[para_idx], rest[2:]

    assert isinstance(block, Paragraph)
    return block, rest


def apply_change(blocks, doc, change_id, new_text):
    paragraph, rest = resolve_paragraph(blocks, doc, change_id.split("/"))
    if len(rest) == 0:
        replace_paragraph_text(paragraph, new_text)
    else:
        set_run_text(paragraph, int(rest[0]), new_text)


def main():
    in_path, changes_path, out_path = sys.argv[1], sys.argv[2], sys.argv[3]
    with open(changes_path, encoding="utf-8") as f:
        changes = json.load(f)

    doc = Document(in_path)
    blocks = list(doc.iter_inner_content())

    applied, errors = 0, []
    for change in changes:
        try:
            apply_change(blocks, doc, change["id"], change["new_text"])
            applied += 1
        except Exception as e:
            errors.append(f"{change['id']!r}: {e}")

    doc.save(out_path)

    print(f"Applied {applied}/{len(changes)} change(s) -> {out_path}")
    if errors:
        print(f"{len(errors)} change(s) failed:", file=sys.stderr)
        for e in errors:
            print(f"  - {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
