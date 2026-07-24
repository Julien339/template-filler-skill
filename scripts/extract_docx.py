"""
Walks an EXISTING .docx and emits a flat, addressable content map of every
text run's current text (body paragraphs, tables, and each section's
DEFAULT header/footer), keyed by a stable structural ID -- the artifact an
agent reads to decide what to change before writing changes.json.

Uses Document.iter_inner_content() (python-docx >=1.2) to walk the body in
true document order -- paragraphs and tables interleaved exactly as they
appear -- instead of python-docx's separate .paragraphs/.tables lists,
which lose that interleaving.

ID scheme:
  body paragraph run:        "{block}/{run}"
  body table cell run:       "{block}/r{row}c{col}/{para}/{run}"
  header/footer paragraph:   "h{section}/{para}/{run}" / "f{section}/{para}/{run}"
An ID with no trailing run index addresses the whole paragraph -- see
apply_docx.py's docstring for what that means on write-back. Only each
section's DEFAULT header/footer is addressed in V1; first-page/even-page
variants are a known limit (see README).

Merged table cells: table.cell(r,c) returns the same underlying XML element
for every coordinate a merge spans -- detected by element identity so only
the origin cell is emitted once.

Runs inside a Word field (date/page-number/TOC auto-fields, built from
w:fldChar begin/separate/end + w:instrText) are flagged "field": true and
should be treated as read-only -- Word silently regenerates their content on
next open, so writing to them looks successful but doesn't stick.

Usage:
    python extract_docx.py <template.docx> <content_map.json>
"""
import sys
import json
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn


def table_origin_cells(table):
    """Yield (row, col, cell) once per distinct cell, skipping the
    duplicate cell references python-docx returns for merged regions.

    Dedup by the underlying lxml element itself (kept alive in `seen`), not
    by id(cell._tc): python-docx builds a fresh, unreferenced _Cell/_tc
    wrapper on every table.cell() call, so an id()-based set with no strong
    reference to the wrapper sees CPython immediately recycle that memory
    address for the next wrapper -- producing spurious "duplicate" hits
    between structurally unrelated cells that just happened to reuse the
    same freed address."""
    seen = set()
    for ri in range(len(table.rows)):
        for ci in range(len(table.columns)):
            cell = table.cell(ri, ci)
            if cell._tc in seen:
                continue
            seen.add(cell._tc)
            yield ri, ci, cell


def paragraph_field_flags(paragraph):
    """One bool per paragraph.runs entry: True if that run is part of a
    Word field (date/page-number/TOC) rather than plain authored text."""
    flags = []
    in_field = False
    for run in paragraph.runs:
        r = run._r
        fld_char = r.find(qn("w:fldChar"))
        instr = r.find(qn("w:instrText"))
        if fld_char is not None:
            fld_type = fld_char.get(qn("w:fldCharType"))
            flags.append(True)
            in_field = fld_type != "end"
            continue
        if instr is not None:
            flags.append(True)
            continue
        flags.append(in_field)
    return flags


def emit_paragraph_runs(paragraph, prefix, out_runs):
    runs = paragraph.runs
    if not runs:
        out_runs.append({"id": f"{prefix}/0", "text": "", "paragraph_text": "", "empty": True})
        return
    field_flags = paragraph_field_flags(paragraph)
    para_text = "".join(r.text for r in runs)
    for ri, (run, is_field) in enumerate(zip(runs, field_flags)):
        entry = {"id": f"{prefix}/{ri}", "text": run.text, "paragraph_text": para_text}
        if is_field:
            entry["field"] = True
        out_runs.append(entry)


def emit_table(table, prefix, out_runs):
    for ri, ci, cell in table_origin_cells(table):
        for pi, para in enumerate(cell.paragraphs):
            emit_paragraph_runs(para, f"{prefix}/r{ri}c{ci}/{pi}", out_runs)


def emit_header_footer(part, label, out_runs):
    if part is None or part.is_linked_to_previous:
        return
    for pi, para in enumerate(part.paragraphs):
        emit_paragraph_runs(para, f"{label}/{pi}", out_runs)


def main():
    in_path, out_path = sys.argv[1], sys.argv[2]
    doc = Document(in_path)

    body_runs = []
    for bi, block in enumerate(doc.iter_inner_content()):
        if isinstance(block, Paragraph):
            emit_paragraph_runs(block, str(bi), body_runs)
        elif isinstance(block, Table):
            emit_table(block, str(bi), body_runs)

    header_footer_runs = []
    for si, section in enumerate(doc.sections):
        emit_header_footer(section.header, f"h{si}", header_footer_runs)
        emit_header_footer(section.footer, f"f{si}", header_footer_runs)

    content_map = {
        "format": "docx",
        "source": in_path,
        "section_count": len(doc.sections),
        "body_runs": body_runs,
        "header_footer_runs": header_footer_runs,
    }
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(content_map, f, ensure_ascii=False, indent=2)

    total = len(body_runs) + len(header_footer_runs)
    print(f"Extracted {total} addressable run(s) ({len(body_runs)} body, "
          f"{len(header_footer_runs)} header/footer) -> {out_path}")


if __name__ == "__main__":
    main()
