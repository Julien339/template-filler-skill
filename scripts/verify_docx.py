"""
Validates a .docx for the class of structural error that Word enforces more
strictly than python-docx (or a naive well-formed-XML check) catches — the
kind that makes Word silently "repair" the file on open and drop content,
with no error raised anywhere in the Python pipeline.

Generic: works on any .docx regardless of how it was produced. Shared
verbatim with the sibling docx-builder skill.

Usage:
    python verify_docx.py <file.docx>

Exits 0 and prints "OK" only if the file is clean. Exits 1 and lists every
problem otherwise. Always run this immediately after apply_docx.py.
"""
import sys
import zipfile
import posixpath
from lxml import etree

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# Elements the OOXML WordprocessingML schema allows at most once among a
# given element's children. Duplicating any of these is exactly the kind of
# error that triggers a silent Word repair-and-drop-content on open.
SINGLETON_CHILDREN = {
    W + "pPr", W + "rPr", W + "tcPr", W + "trPr", W + "tblPr",
    W + "sectPr", W + "shd", W + "tblGrid",
}

DOC_PARTS_PREFIXES = ("word/document.xml", "word/header", "word/footer")


def check_zip_integrity(z, problems):
    bad = z.testzip()
    if bad:
        problems.append(f"zip integrity: corrupt entry {bad}")


def check_xml_wellformed(z, problems):
    parts = {}
    for name in z.namelist():
        if name.endswith(".xml") or name.endswith(".rels"):
            try:
                parts[name] = etree.fromstring(z.read(name))
            except Exception as e:
                problems.append(f"malformed XML: {name}: {e}")
    return parts


def is_target_part(name):
    return name == "word/document.xml" or (
        name.startswith(("word/header", "word/footer")) and name.endswith(".xml")
    )


def check_duplicate_singletons(parts, problems):
    for name, root in parts.items():
        if not is_target_part(name):
            continue
        for el in root.iter():
            counts = {}
            for child in el:
                counts[child.tag] = counts.get(child.tag, 0) + 1
            for tag, count in counts.items():
                if tag in SINGLETON_CHILDREN and count > 1:
                    problems.append(
                        f"{name}: <{etree.QName(el).localname}> has {count} "
                        f"<{etree.QName(tag).localname}> children (schema allows at most 1)"
                    )


def check_illegal_chars(parts, problems):
    for name, root in parts.items():
        if not is_target_part(name):
            continue
        text = etree.tostring(root).decode("utf-8", errors="replace")
        for ch in text:
            if ord(ch) < 0x20 and ch not in "\t\n\r":
                problems.append(f"{name}: illegal XML control character U+{ord(ch):04X}")
                break


def check_media_rels(z, problems):
    names = set(z.namelist())
    for name in names:
        parts = name.split("/")
        if "_rels" not in parts or not name.endswith(".rels"):
            continue
        # A .rels part always lives at <base>/_rels/<partname>.rels and
        # describes relationships *from* the part at <base>/<partname> — so
        # relative targets (e.g. "../customXml/item1.xml") resolve against
        # <base>/, not against the .rels file's own directory.
        base_dir = "/".join(parts[:parts.index("_rels")])
        root = etree.fromstring(z.read(name))
        for rel in root:
            target = rel.get("Target")
            if not target or rel.get("TargetMode") == "External":
                continue
            resolved = posixpath.normpath(posixpath.join(base_dir, target)) if base_dir else posixpath.normpath(target)
            if resolved not in names:
                problems.append(f"{name}: dangling relationship target {target!r}")


def check_opens_with_python_docx(path, problems):
    try:
        from docx import Document
        doc = Document(path)
        if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
            problems.append("document has no paragraphs and no tables")
    except Exception as e:
        problems.append(f"python-docx failed to open the file: {e}")


def main():
    path = sys.argv[1]
    problems = []

    with zipfile.ZipFile(path) as z:
        check_zip_integrity(z, problems)
        parts = check_xml_wellformed(z, problems)
        if parts:
            check_duplicate_singletons(parts, problems)
            check_illegal_chars(parts, problems)
        check_media_rels(z, problems)

    check_opens_with_python_docx(path, problems)

    if problems:
        print(f"FAIL — {len(problems)} problem(s) in {path}:", file=sys.stderr)
        for p in problems:
            print(f"  - {p}", file=sys.stderr)
        sys.exit(1)

    print(f"OK — {path} passed structural validation.")


if __name__ == "__main__":
    main()
