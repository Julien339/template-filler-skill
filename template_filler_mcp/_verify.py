"""Structural validation for PPTX and DOCX files.

Catches the class of XML corruption that PowerPoint/Word enforce more strictly
than python-pptx/python-docx — the kind that makes the application silently
"repair" the file on open and drop content, with no error raised in Python.
"""

from __future__ import annotations

import zipfile
from typing import Any

from lxml import etree

from ._path_validation import validate_input_file

# ── OOXML namespaces ───────────────────────────────────────────────────

A = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# Elements the schema allows at most once among a given element's children.
PPTX_SINGLETON_CHILDREN = {
    A + "effectLst",
    A + "effectDag",
    A + "xfrm",
    A + "ln",
    A + "scene3d",
    A + "sp3d",
    A + "custGeom",
    A + "prstGeom",
}

DOCX_SINGLETON_CHILDREN = {
    W + "pPr",
    W + "rPr",
    W + "tcPr",
    W + "trPr",
    W + "tblPr",
    W + "sectPr",
    W + "shd",
    W + "tblGrid",
}


# ── shared checks ──────────────────────────────────────────────────────


def _check_zip_integrity(z, problems):
    bad = z.testzip()
    if bad:
        problems.append(f"zip integrity: corrupt entry {bad}")


def _check_xml_wellformed(z, problems):
    parts = {}
    for name in z.namelist():
        if name.endswith(".xml") or name.endswith(".rels"):
            try:
                parts[name] = etree.fromstring(z.read(name))
            except Exception as e:
                problems.append(f"malformed XML: {name}: {e}")
    return parts


def _check_duplicate_singletons(parts, problems, target_filter, singleton_set):
    for name, root in parts.items():
        if not target_filter(name):
            continue
        for el in root.iter():
            counts = {}
            for child in el:
                counts[child.tag] = counts.get(child.tag, 0) + 1
            for tag, count in counts.items():
                if tag in singleton_set and count > 1:
                    problems.append(
                        f"{name}: <{etree.QName(el).localname}> has {count} "
                        f"<{etree.QName(tag).localname}> children (schema allows at most 1)"
                    )


def _check_illegal_chars(parts, problems, target_filter):
    for name, root in parts.items():
        if not target_filter(name):
            continue
        text = etree.tostring(root).decode("utf-8", errors="replace")
        for ch in text:
            if ord(ch) < 0x20 and ch not in "\t\n\r":
                problems.append(f"{name}: illegal XML control character U+{ord(ch):04X}")
                break


# ── PPTX verification ──────────────────────────────────────────────────


def _pptx_target_filter(name):
    return name.startswith("ppt/slides/slide") and name.endswith(".xml")


def _check_media_rels_pptx(z, problems):
    names = set(z.namelist())
    for name in names:
        if not name.startswith("ppt/slides/_rels/"):
            continue
        root = etree.fromstring(z.read(name))
        for rel in root:
            target = rel.get("Target")
            if not target or rel.get("TargetMode") == "External":
                continue
            resolved = "ppt/" + target[3:] if target.startswith("../") else target
            if resolved not in names:
                problems.append(f"{name}: dangling relationship target {target!r}")


def _check_opens_with_python_pptx(path, problems):
    try:
        from pptx import Presentation

        prs = Presentation(path)
        if len(prs.slides) == 0:
            problems.append("presentation has 0 slides")
    except Exception as e:
        problems.append(f"python-pptx failed to open the file: {e}")


def verify_pptx_structure(filepath: str) -> dict[str, Any]:
    """Validate a PPTX file for structural/XML integrity.

    Returns {"ok": True, "problems": []} if the file is clean,
    or {"ok": False, "problems": [...]} with a list of issues.
    """
    validated_path = validate_input_file(filepath, allowed_extensions=(".pptx",))

    problems = []
    try:
        z = zipfile.ZipFile(str(validated_path))
    except (zipfile.BadZipFile, FileNotFoundError) as e:
        return {"ok": False, "problems": [f"cannot open as ZIP: {e}"]}
    with z:
        _check_zip_integrity(z, problems)
        parts = _check_xml_wellformed(z, problems)
        if parts:
            _check_duplicate_singletons(parts, problems, _pptx_target_filter, PPTX_SINGLETON_CHILDREN)
            _check_illegal_chars(parts, problems, _pptx_target_filter)
        _check_media_rels_pptx(z, problems)
    _check_opens_with_python_pptx(str(validated_path), problems)

    return {"ok": len(problems) == 0, "problems": problems}


# ── DOCX verification ──────────────────────────────────────────────────


def _docx_target_filter(name):
    return name == "word/document.xml" or (
        name.startswith(("word/header", "word/footer")) and name.endswith(".xml")
    )


def _check_media_rels_docx(z, problems):
    import posixpath

    names = set(z.namelist())
    for name in names:
        parts = name.split("/")
        if "_rels" not in parts or not name.endswith(".rels"):
            continue
        base_dir = "/".join(parts[: parts.index("_rels")])
        root = etree.fromstring(z.read(name))
        for rel in root:
            target = rel.get("Target")
            if not target or rel.get("TargetMode") == "External":
                continue
            resolved = (
                posixpath.normpath(posixpath.join(base_dir, target))
                if base_dir
                else posixpath.normpath(target)
            )
            if resolved not in names:
                problems.append(f"{name}: dangling relationship target {target!r}")


def _check_opens_with_python_docx(path, problems):
    try:
        from docx import Document

        doc = Document(path)
        if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
            problems.append("document has no paragraphs and no tables")
    except Exception as e:
        problems.append(f"python-docx failed to open the file: {e}")


def verify_docx_structure(filepath: str) -> dict[str, Any]:
    """Validate a DOCX file for structural/XML integrity.

    Returns {"ok": True, "problems": []} if the file is clean,
    or {"ok": False, "problems": [...]} with a list of issues.
    """
    validated_path = validate_input_file(filepath, allowed_extensions=(".pptx",))

    problems = []
    try:
        z = zipfile.ZipFile(str(validated_path))
    except (zipfile.BadZipFile, FileNotFoundError) as e:
        return {"ok": False, "problems": [f"cannot open as ZIP: {e}"]}
    with z:
        _check_zip_integrity(z, problems)
        parts = _check_xml_wellformed(z, problems)
        if parts:
            _check_duplicate_singletons(parts, problems, _docx_target_filter, DOCX_SINGLETON_CHILDREN)
            _check_illegal_chars(parts, problems, _docx_target_filter)
        _check_media_rels_docx(z, problems)
    _check_opens_with_python_docx(str(validated_path), problems)

    return {"ok": len(problems) == 0, "problems": problems}
